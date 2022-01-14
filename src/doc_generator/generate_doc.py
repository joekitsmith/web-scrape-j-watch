from typing import Dict, List

from docx import Document
from docx.shared import Cm, Pt
from docx.shared import RGBColor
from docx.oxml.shared import qn
from docx.oxml.xmlchemy import OxmlElement
from src.doc_generator.clean_text import clean_text, remove_open_in_new_tab
from src.web_scraper.data_classes import ContentKeys

from web_scraper.extract_content import Line


class DocumentGenerator:
    def __init__(self) -> None:
        self.document = Document()

        self._configure_document()

    def add_page(self, data: Dict[str, List[Line]]) -> None:
        """
        Add all article content to document.

        Arguments
        ---------
        data: {str, list}
            keys are sections of article
            values are lists containing Line objects
        """
        self.data = data

        self._create_heading()
        self._create_contributors()
        self._create_summary()
        self._create_main_content()
        self._create_comment()

        self.document.add_page_break()

    def _configure_document(self) -> None:
        """
        Set document margins.
        """
        margin = 2.4
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(margin)
            section.bottom_margin = Cm(margin)
            section.left_margin = Cm(margin)
            section.right_margin = Cm(margin)

    def _create_heading(self) -> None:
        """
        Add article title to document.
        """
        text = clean_text(self.data[ContentKeys.TITLE])
        run = self.document.add_paragraph().add_run(text)
        font = run.font
        font.name = "Arial"
        font.size = Pt(24)
        font.bold = True
        font.color.rgb = RGBColor(99, 62, 106)

    def _create_contributors(self):
        """
        Add contributors section to document.
        """
        text = clean_text(self.data[ContentKeys.CONTRIBUTORS])
        if text:
            run = self.document.add_paragraph().add_run(text)
            font = run.font
            font.name = "Times New Roman"
            font.size = Pt(9)
            font.italic = True

    def _create_summary(self):
        """
        Add summary section to document.
        """
        text = clean_text(self.data[ContentKeys.SUMMARY])
        if text:
            run = self.document.add_paragraph().add_run(text)
            font = run.font
            font.name = "Arial"
            font.size = Pt(9)
            font.italic = True
            font.color.rgb = RGBColor(0, 0, 0)

    def _create_main_content(self):
        """
        Add main content to document.
        """
        for line in remove_open_in_new_tab(self.data[ContentKeys.MAIN_CONTENT]):
            if line.marker:
                # create bullet point list
                paragraph = self.document.add_paragraph(style="List Bullet")
            else:
                paragraph = self.document.add_paragraph()
            run = paragraph.add_run(line.text)
            font = run.font
            font.name = "Arial"
            font.size = Pt(9)
            font.color.rgb = RGBColor(0, 0, 0)

    def _create_comment(self):
        """
        Add comment section to document.
        """
        header_text = clean_text(self.data[ContentKeys.COMMENT_HEADER])
        content_text = clean_text(self.data[ContentKeys.COMMENT_CONTENT])

        if header_text:
            # use table to add highlighted section
            table = self.document.add_table(rows=2, cols=1)
            header_cell = table.rows[0].cells[0]
            body_cell = table.rows[1].cells[0]

            header_cell.text = header_text
            body_cell.text = content_text

            # comment section is present
            for i in range(2):
                self._set_cell_background(table.rows[i].cells[0], "E3E9F0")

            for i, row in enumerate(table.rows):
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            # header
                            if i == 0:
                                font.size = Pt(14)
                                font.bold = True
                            # body
                            if i == 1:
                                font.name = "Arial"
                                font.size = Pt(9)

    def _set_cell_background(self, cell, fill: str):
        """
        Set background of cell in table to defined colour.

        Arguments
        ---------
        cell :
            table cell to fill
        fill : str
            color to be used for the background
        """

        cell_properties = cell._element.tcPr

        cell_shading = OxmlElement("w:shd")  # add new w:shd element to it

        cell_shading.set(qn("w:fill"), fill)  # set fill property, respecting namespace
        cell_properties.append(
            cell_shading
        )  # finally extend cell props with shading element
